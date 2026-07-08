#!/usr/bin/env python3
"""Generate a BBMS report as DOCX using python-docx.

Usage: python3 generate_report_docx.py <json_file> <output_docx>
"""
import sys
import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Colors (RGBColor)
ROSE = RGBColor(0xe1, 0x1d, 0x48)
DARK = RGBColor(0x1e, 0x29, 0x3b)
MUTED = RGBColor(0x64, 0x74, 0x8b)
EMERALD = RGBColor(0x10, 0xb9, 0x81)
AMBER = RGBColor(0xf5, 0x9e, 0x0b)
SKY = RGBColor(0x0e, 0xa5, 0xe9)
VIOLET = RGBColor(0x8b, 0xb5, 0xc6)
WHITE = RGBColor(0xff, 0xff, 0xff)


def shade_cell(cell, hex_color):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9, align='left'):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=None):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = DARK
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = color or ROSE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_meta_line(doc, text):
    """Add a muted meta paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run.font.italic = True
    p.paragraph_format.space_after = Pt(12)


def build_stat_table(doc, summary):
    """4-column summary table."""
    items = [
        ('Total Units', summary.get('totalUnits', 0)),
        ('Available', summary.get('available', 0)),
        ('Reserved', summary.get('reserved', 0)),
        ('Issued', summary.get('issued', 0)),
        ('Expired', summary.get('expired', 0)),
        ('Discarded', summary.get('discarded', 0)),
        ('Expiring Soon', summary.get('nearExpiry', 0)),
        ('Total Activity', summary.get('totalUnits', 0)),
    ]
    # Build 2 rows of 4 columns
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    for i, (label, value) in enumerate(items):
        row = i // 4
        col = i % 4
        cell = table.rows[row].cells[col]
        cell.text = ''
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.size = Pt(8)
        r1.font.color.rgb = MUTED
        r1.font.bold = True
        p2 = cell.add_paragraph()
        r2 = p2.add_run(str(value))
        r2.font.size = Pt(18)
        r2.font.bold = True
        r2.font.color.rgb = DARK


def build_blood_group_table(doc, stock_by_group):
    """Detailed blood group stock table."""
    headers = ['Blood Group', 'Available', 'Total', 'Expired', 'Reserved', 'Issued']
    table = doc.add_table(rows=1 + len(stock_by_group), cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=WHITE, size=10, align='center' if i > 0 else 'left')
        shade_cell(cell, 'e11d48')

    # Data rows
    for r, g in enumerate(sorted(stock_by_group.keys())):
        v = stock_by_group[g]
        row = table.rows[r + 1]
        set_cell_text(row.cells[0], g, bold=True, color=WHITE, size=10, align='center')
        shade_cell(row.cells[0], 'ef4444' if '+' in g and 'AB' not in g else
                              'dc2626' if 'O-' == g else
                              '10b981' if 'A+' == g else
                              '059669' if 'A-' == g else
                              '0ea5e9' if 'B+' == g else
                              '0284c7' if 'B-' == g else
                              '8b5cf6' if 'AB+' == g else
                              '7c3aed')
        set_cell_text(row.cells[1], v.get('available', 0), size=10, align='center')
        set_cell_text(row.cells[2], v.get('total', 0), size=10, align='center')
        set_cell_text(row.cells[3], v.get('expired', 0), size=10, align='center')
        set_cell_text(row.cells[4], v.get('reserved', 0), size=10, align='center')
        set_cell_text(row.cells[5], v.get('issued', 0), size=10, align='center')


def build_status_table(doc, status_dict, title):
    """Status breakdown table."""
    if not status_dict:
        doc.add_paragraph('No data')
        return
    items = sorted(status_dict.items(), key=lambda x: -x[1])
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = 'Light Grid Accent 1'
    set_cell_text(table.rows[0].cells[0], title, bold=True, color=WHITE, size=10)
    set_cell_text(table.rows[0].cells[1], 'Count', bold=True, color=WHITE, size=10, align='center')
    shade_cell(table.rows[0].cells[0], '1e293b')
    shade_cell(table.rows[0].cells[1], '1e293b')
    for i, (k, v) in enumerate(items):
        set_cell_text(table.rows[i + 1].cells[0], k, size=10)
        set_cell_text(table.rows[i + 1].cells[1], v, size=10, align='center', bold=True)


def build_storage_table(doc, storage):
    """Storage utilization table."""
    if not storage:
        doc.add_paragraph('No storage units')
        return
    headers = ['Storage Unit', 'Category', 'Used', 'Capacity', 'Utilization']
    table = doc.add_table(rows=1 + len(storage), cols=5)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=WHITE, size=10, align='center' if i > 1 else 'left')
        shade_cell(cell, '8b5cf6')
    for r, s in enumerate(storage):
        row = table.rows[r + 1]
        set_cell_text(row.cells[0], s.get('name', ''), size=9)
        set_cell_text(row.cells[1], s.get('category', ''), size=9)
        set_cell_text(row.cells[2], s.get('used', 0), size=10, align='center')
        set_cell_text(row.cells[3], s.get('capacity', 0), size=10, align='center')
        util = s.get('utilization', 0)
        util_color = ROSE if util > 90 else AMBER if util > 70 else EMERALD
        set_cell_text(row.cells[4], f'{util}%', size=10, align='center', bold=True, color=util_color)


def build_activity_table(doc, activity):
    """Recent activity table."""
    if not activity:
        doc.add_paragraph('No recent activity')
        return
    headers = ['Timestamp', 'Action', 'Description', 'User']
    table = doc.add_table(rows=1 + min(15, len(activity)), cols=4)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=WHITE, size=9)
        shade_cell(cell, '1e293b')
    for i, a in enumerate(activity[:15]):
        ts = a.get('createdAt', '')[:16].replace('T', ' ')
        row = table.rows[i + 1]
        set_cell_text(row.cells[0], ts, size=8)
        set_cell_text(row.cells[1], a.get('action', ''), size=8)
        set_cell_text(row.cells[2], a.get('description', '')[:60], size=8)
        set_cell_text(row.cells[3], a.get('user', '')[:25], size=8)


def main():
    if len(sys.argv) < 3:
        print('Usage: generate_report_docx.py <json_file> <output_docx>', file=sys.stderr)
        sys.exit(1)

    json_file = sys.argv[1]
    output_docx = sys.argv[2]

    with open(json_file) as f:
        data = json.load(f)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    add_heading(doc, 'Blood Bank Management Report', level=0)

    # Subtitle (facility + period)
    facility = data.get('facility') or {}
    facility_name = facility.get('name', 'All Facilities') if isinstance(facility, dict) else 'All Facilities'
    facility_type = facility.get('type', '') if isinstance(facility, dict) else ''
    facility_region = facility.get('region', '') if isinstance(facility, dict) else ''
    period = data.get('period') or {}
    if period.get('from') or period.get('to'):
        period_str = f"Period: {period.get('from', 'start')} to {period.get('to', 'now')}"
    else:
        period_str = 'Period: All time'
    subtitle = facility_name
    if facility_type:
        subtitle += f" · {facility_type}"
    if facility_region:
        subtitle += f" · {facility_region} Region"
    add_meta_line(doc, f"{subtitle} · {period_str}")

    # Summary
    add_heading(doc, 'Inventory Summary', level=1)
    build_stat_table(doc, data.get('summary', {}))

    # Blood group table
    add_heading(doc, 'Stock by Blood Group', level=1)
    build_blood_group_table(doc, data.get('stockByGroup', {}))

    # Component distribution
    add_heading(doc, 'Component Distribution (Available Units)', level=1)
    comp = data.get('componentDist', {})
    if comp:
        table = doc.add_table(rows=1 + len(comp), cols=2)
        table.style = 'Light Grid Accent 1'
        set_cell_text(table.rows[0].cells[0], 'Component', bold=True, color=WHITE, size=10)
        set_cell_text(table.rows[0].cells[1], 'Count', bold=True, color=WHITE, size=10, align='center')
        shade_cell(table.rows[0].cells[0], 'e11d48')
        shade_cell(table.rows[0].cells[1], 'e11d48')
        for i, (k, v) in enumerate(sorted(comp.items(), key=lambda x: -x[1])):
            set_cell_text(table.rows[i + 1].cells[0], k, size=10)
            set_cell_text(table.rows[i + 1].cells[1], v, size=10, align='center', bold=True)

    doc.add_page_break()

    # Internal requests
    add_heading(doc, 'Internal Blood Requests', level=1)
    ir = data.get('internalRequests', {})
    p = doc.add_paragraph()
    p.add_run(f'Total requests: ').font.bold = True
    p.add_run(str(ir.get('total', 0))).font.size = Pt(11)
    doc.add_paragraph()
    build_status_table(doc, ir.get('byStatus', {}), 'Status Breakdown')
    doc.add_paragraph()
    build_status_table(doc, ir.get('byUrgency', {}), 'Urgency Breakdown')

    # Network requests
    add_heading(doc, 'Network Blood Requests (Broadcast)', level=1)
    nr = data.get('networkRequests', {})
    p = doc.add_paragraph()
    p.add_run(f'Total network requests: ').font.bold = True
    p.add_run(str(nr.get('total', 0))).font.size = Pt(11)
    doc.add_paragraph()
    build_status_table(doc, nr.get('byStatus', {}), 'Status Breakdown')
    doc.add_paragraph()
    build_status_table(doc, nr.get('byGroup', {}), 'Blood Group Breakdown')

    doc.add_page_break()

    # Storage
    add_heading(doc, 'Storage Utilization', level=1)
    build_storage_table(doc, data.get('storageUtilization', []))

    # Recent activity
    add_heading(doc, 'Recent Activity (Audit Log)', level=1)
    build_activity_table(doc, data.get('recentActivity', []))

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'Generated by BBMS Ghana on {datetime.now().strftime("%d %B %Y at %H:%M")}')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = MUTED

    doc.save(output_docx)
    print(f'DOCX generated: {output_docx}')


if __name__ == '__main__':
    main()
